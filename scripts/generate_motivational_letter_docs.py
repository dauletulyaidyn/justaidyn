from __future__ import annotations

from pathlib import Path
from typing import Iterable

from deep_translator import GoogleTranslator
from docx import Document


SOURCE_DOC = Path(
    r"C:\Users\JustAidyn\Desktop\Motivational letter\Applied_AI_Agents_Online_Course_Agreement_KK_2_versions_v11_marked_changes.docx"
)
OUTPUT_DIR = Path("downloads/motivational-letter")
KK_OUTPUT = OUTPUT_DIR / "Applied_AI_Agents_Online_Course_Agreement_KK.docx"
RU_OUTPUT = OUTPUT_DIR / "Applied_AI_Agents_Online_Course_Agreement_RU.docx"


def chunk_text(text: str, max_length: int = 3500) -> Iterable[str]:
    if len(text) <= max_length:
        yield text
        return

    words = text.split(" ")
    current = []
    current_len = 0

    for word in words:
        extra = len(word) + (1 if current else 0)
        if current and current_len + extra > max_length:
            yield " ".join(current)
            current = [word]
            current_len = len(word)
        else:
            current.append(word)
            current_len += extra

    if current:
        yield " ".join(current)


def translate_to_ru(text: str, translator: GoogleTranslator) -> str:
    raw = text.strip()
    if not raw:
        return text
    if set(raw) == {"_"}:
        return text

    parts = []
    for chunk in chunk_text(raw):
        translated = translator.translate(chunk)
        parts.append(translated if translated else chunk)

    return " ".join(parts)


def build_ru_doc(kk_doc: Document) -> Document:
    ru_doc = Document()
    translator = GoogleTranslator(source="auto", target="ru")

    for paragraph in kk_doc.paragraphs:
        translated_text = translate_to_ru(paragraph.text, translator)
        new_paragraph = ru_doc.add_paragraph(translated_text)
        if paragraph.style:
            try:
                new_paragraph.style = paragraph.style
            except Exception:
                pass

    return ru_doc


def main() -> None:
    if not SOURCE_DOC.exists():
        raise FileNotFoundError(f"Source document not found: {SOURCE_DOC}")

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    kk_doc = Document(SOURCE_DOC)

    kk_doc.save(KK_OUTPUT)
    ru_doc = build_ru_doc(kk_doc)
    ru_doc.save(RU_OUTPUT)

    print(f"Saved: {KK_OUTPUT.resolve()}")
    print(f"Saved: {RU_OUTPUT.resolve()}")


if __name__ == "__main__":
    main()
