from pathlib import Path

from docx import Document


BASE_DIR = Path("downloads/motivational-letter")

KK_FULL = BASE_DIR / "Applied_AI_Agents_Online_Course_Agreement_KK.docx"
RU_FULL = BASE_DIR / "Applied_AI_Agents_Online_Course_Agreement_RU.docx"

KK_18 = BASE_DIR / "Applied_AI_Agents_Online_Course_Agreement_KK_18_plus.docx"
KK_16_17 = BASE_DIR / "Applied_AI_Agents_Online_Course_Agreement_KK_16_17.docx"
RU_18 = BASE_DIR / "Applied_AI_Agents_Online_Course_Agreement_RU_18_plus.docx"
RU_16_17 = BASE_DIR / "Applied_AI_Agents_Online_Course_Agreement_RU_16_17.docx"


def build_doc_from_paragraphs(source: Document, paragraphs: list) -> Document:
    target = Document()
    for paragraph in paragraphs:
        new_paragraph = target.add_paragraph(paragraph.text)
        if paragraph.style:
            try:
                new_paragraph.style = paragraph.style
            except Exception:
                pass
    return target


def split_doc(
    source_path: Path,
    marker_adult_variants: tuple[str, ...],
    marker_teen_variants: tuple[str, ...],
    adult_output: Path,
    teen_output: Path,
) -> None:
    doc = Document(source_path)
    paragraphs = doc.paragraphs

    adult_idx = None
    teen_idx = None

    for index, paragraph in enumerate(paragraphs):
        text = paragraph.text.strip()
        if adult_idx is None and any(marker in text for marker in marker_adult_variants):
            adult_idx = index
        if teen_idx is None and any(marker in text for marker in marker_teen_variants):
            teen_idx = index

    if teen_idx is None:
        raise ValueError(f"Teen section marker not found in {source_path.name}")

    if adult_idx is None:
        adult_idx = 0

    adult_doc = build_doc_from_paragraphs(doc, paragraphs[adult_idx:teen_idx])
    teen_doc = build_doc_from_paragraphs(doc, paragraphs[teen_idx:])

    adult_doc.save(adult_output)
    teen_doc.save(teen_output)

    print(f"Saved: {adult_output.resolve()}")
    print(f"Saved: {teen_output.resolve()}")


def main() -> None:
    if not KK_FULL.exists() or not RU_FULL.exists():
        raise FileNotFoundError("Full KK/RU agreement docs are missing in downloads/motivational-letter")

    split_doc(
        source_path=KK_FULL,
        marker_adult_variants=("Нұсқа A", "Нұсқа А"),
        marker_teen_variants=("Нұсқа B", "Нұсқа В"),
        adult_output=KK_18,
        teen_output=KK_16_17,
    )

    split_doc(
        source_path=RU_FULL,
        marker_adult_variants=("Вариант A", "Вариант А"),
        marker_teen_variants=("Вариант B", "Вариант В", "Несовершеннолетний участник (16-17 лет)"),
        adult_output=RU_18,
        teen_output=RU_16_17,
    )


if __name__ == "__main__":
    main()
